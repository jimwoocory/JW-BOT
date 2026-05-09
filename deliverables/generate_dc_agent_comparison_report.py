from __future__ import annotations

from collections import Counter, defaultdict
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENTATION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from openpyxl import load_workbook


ROOT = Path("/Users/dianchi/DC-Agent")
XLSX = Path("/Users/dianchi/Desktop/文档/问卷.xlsx")
OUT_DOCX = ROOT / "deliverables" / "DC-Agent二期需求对比报告.docx"

FONT = "Microsoft YaHei"


def split_multi(value: object) -> list[str]:
    if not value:
        return []
    return [item.strip() for item in str(value).split(",") if item.strip()]


def set_run_font(
    run, size: float | None = None, bold: bool | None = None, color: str | None = None
):
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def set_paragraph_font(
    paragraph, size: float = 16, bold: bool | None = None, color: str | None = None
):
    for run in paragraph.runs:
        set_run_font(run, size=size, bold=bold, color=color)


def shade_cell(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(
    cell, top: int = 65, start: int = 120, bottom: int = 65, end: int = 120
):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_cell_text(
    cell, text: str, size: float = 12.9, bold: bool = False, color: str = "1D2939"
):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.03
    r = p.add_run(text)
    set_run_font(r, size=size, bold=bold, color=color)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_margins(cell)


def add_table(
    doc: Document,
    headers: list[str],
    rows: list[list[str]],
    widths: list[float] | None = None,
):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        set_cell_text(hdr[i], h, size=13.8, bold=True, color="073763")
        shade_cell(hdr[i], "EAF3FF")
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            set_cell_text(cells[i], val, size=12.9)
    if widths:
        for row in table.rows:
            for i, width in enumerate(widths):
                row.cells[i].width = Cm(width)
    doc.add_paragraph()
    return table


def add_heading(doc: Document, text: str, level: int = 1):
    p = doc.add_paragraph()
    if level == 1:
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(6)
        size = 23
        color = "073763"
    else:
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(5)
        size = 18
        color = "174A80"
    r = p.add_run(text)
    set_run_font(r, size=size, bold=True, color=color)
    return p


def add_body(
    doc: Document, text: str, size: float = 16.5, bold_prefix: str | None = None
):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(7)
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_run_font(r1, size=size, bold=True, color="073763")
        r2 = p.add_run(text[len(bold_prefix) :])
        set_run_font(r2, size=size)
    else:
        r = p.add_run(text)
        set_run_font(r, size=size)
    return p


def add_bullets(doc: Document, items: list[str]):
    for item in items:
        p = doc.add_paragraph(style=None)
        p.paragraph_format.left_indent = Cm(0.55)
        p.paragraph_format.first_line_indent = Cm(-0.25)
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run("• ")
        set_run_font(r, size=16.5, color="0B57A4")
        r = p.add_run(item)
        set_run_font(r, size=16.5)


def add_callout(doc: Document, text: str, fill: str = "F3F8FF"):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    shade_cell(cell, fill)
    set_cell_text(cell, text, size=16.2, bold=False)
    doc.add_paragraph()


def start_page(doc: Document, title: str):
    doc.add_page_break()
    add_heading(doc, title, 1)


def load_data():
    ws = load_workbook(XLSX, data_only=True).active
    rows = list(ws.iter_rows(values_only=True))
    headers = list(rows[0])
    return [dict(zip(headers, row)) for row in rows[1:]]


def pct(count: int, total: int) -> str:
    return f"{count}/{total}（{round(count / total * 100)}%）"


def build_report():
    data = load_data()
    total = len(data)
    dept_counter = Counter(row["你的岗位/部门是？"] for row in data)
    dept_people: dict[str, list[str]] = defaultdict(list)
    for row in data:
        dept_people[row["你的岗位/部门是？"]].append(row["提交人"])

    def count_col(col: str) -> Counter[str]:
        counter: Counter[str] = Counter()
        for row in data:
            counter.update(split_multi(row[col]))
        return counter

    first_phase = count_col("如果第一阶段只能优先上线 3 个功能，你会选哪 3 个？")
    work_need = count_col("你平时最希望 DC-Agent 帮你处理哪类工作？")
    usage = count_col("你希望通过什么方式使用 DC-Agent？")
    data_sources = count_col("哪些数据你希望 DC-Agent 能读取或对接？")
    auto_reply = count_col("你能接受 DC-Agent 在群聊中自动回应吗？")
    content_need = count_col("关于内容创作，你最希望 DC-Agent 帮你生成什么？")

    doc = Document()
    sec = doc.sections[0]
    sec.orientation = WD_ORIENTATION.LANDSCAPE
    sec.page_width = Cm(29.7)
    sec.page_height = Cm(21.0)
    sec.top_margin = Cm(0.95)
    sec.bottom_margin = Cm(0.95)
    sec.left_margin = Cm(1.15)
    sec.right_margin = Cm(1.15)
    sec.header_distance = Cm(0)
    sec.footer_distance = Cm(0)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    normal.font.size = Pt(16.5)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run("DC-Agent 二期需求对比报告")
    set_run_font(r, size=31, bold=True, color="073763")
    p = doc.add_paragraph()
    r = p.add_run(
        "基于飞书问卷答卷统计，对比 Claude Code 已形成的二期开发方案，明确 DC-Agent 二期交付顺序、功能缺口和调整路线。"
    )
    set_run_font(r, size=17, color="344054")

    add_table(
        doc,
        ["系统名称", "答卷数量", "参与部门", "报告日期"],
        [
            [
                "DC-Agent",
                f"{total} 份有效答卷",
                "品宣部、中台运营、执行运营、财务部、其他",
                "2026-05-08",
            ]
        ],
        [4.0, 4.2, 12.5, 4.6],
    )

    add_callout(
        doc,
        "总判断：Claude Code 已做好的二期方案偏底层架构升级，重点是 Router、Harness、知识库、记忆和 Hermes 深度执行闭环；本次问卷反馈偏员工可感知功能，重点是任务提醒、项目群聊总结、资料查询和方案生成。两者方向不冲突，但二期交付顺序需要调整：先把员工高频痛点做出来，再逐步接上底层深度执行能力。",
    )

    start_page(doc, "一、本次答卷样本统计")
    add_body(
        doc,
        f"本次共收集 {total} 份有效答卷，提交时间集中在 2026-05-06。答卷覆盖 5 类部门/岗位，其中品宣部和中台运营各 4 人，是本次需求反馈的主要来源。",
    )
    dept_rows = []
    for dept, count in dept_counter.most_common():
        dept_rows.append([dept, pct(count, total), "、".join(dept_people[dept])])
    add_table(doc, ["部门/岗位", "人数占比", "答卷提交人"], dept_rows, [5.0, 5.0, 16.0])

    start_page(doc, "二、问卷需求统计汇总")
    add_table(
        doc,
        ["需求项", "选择人数", "开发含义"],
        [
            [
                "任务提取和提醒",
                pct(first_phase["任务提取和提醒"], total),
                "员工最强需求，必须作为 DC-Agent 二期 P0。",
            ],
            [
                "项目群聊总结",
                pct(first_phase["项目群聊总结"], total),
                "项目群信息分散是主要痛点，需要能总结客户反馈、老板要求、执行问题和阶段结论。",
            ],
            [
                "飞书文档读取",
                pct(data_sources["飞书文档"], total),
                "资料查询应优先接飞书文档，先做白名单资料源。",
            ],
            [
                "飞书表格读取",
                pct(data_sources["飞书表格"], total),
                "任务台账、项目资料、客户资料可优先落到飞书表格。",
            ],
            [
                "私聊查询或生成",
                pct(usage["私聊 DC-Agent 查询或生成内容"], total),
                "员工更偏低打扰使用，不建议默认群聊自动抢答。",
            ],
            [
                "客户提案大纲",
                pct(content_need["客户提案大纲"], total),
                "内容/方案生成仍是核心场景，但要结合客户资料和项目上下文。",
            ],
        ],
        [6.0, 4.5, 15.5],
    )

    start_page(doc, "三、参与部门的真实需求画像")
    add_table(
        doc,
        ["部门", "主要需求", "代表性开放反馈"],
        [
            [
                "品宣部",
                "任务提醒、项目总结、内容重点整理、社媒数据采集、私域/竞品负面监控。",
                "收集小红书、抖音等社媒平台的用户发帖链接，并填写互动量、传播量；监控私域社群本竞品负面信息。",
            ],
            [
                "中台运营",
                "客户资料、项目资料、历史方案查询；项目群总结；客户方案框架和应标文件。",
                "快速调取公司现有网盘及飞书存档资料链接；根据客户需求准确生成方案框架；帮忙编应标文件。",
            ],
            ["执行运营", "用户反馈信息提取和总结。", "提取用户反馈信息总结。"],
            [
                "财务部",
                "银行流水、发票、客户供应商应收应付、报销单自动匹配对账。",
                "银行流水与发票和客户供应商应收应付，报销单自动匹配对账。",
            ],
            [
                "其他",
                "工作安排、进度实时更新、图片生成。",
                "人工工作安排和进度实时更新；能够生成图片。",
            ],
        ],
        [4.0, 10.0, 12.0],
    )

    start_page(doc, "四、Claude Code 已形成的二期开发方案摘要")
    add_table(
        doc,
        ["模块", "Claude Code 方案重点", "价值判断"],
        [
            [
                "AstrBot 前台",
                "员工入口、QQ/飞书接入、即时对话、结果推送。",
                "正确。DC-Agent 不应让 Hermes 直接面对员工，AstrBot 应继续作为前台。",
            ],
            [
                "Router",
                "识别普通对话、技能调用、正式任务、升级信号。",
                "正确。后续任务总结、资料查询、任务提醒都需要 Router 分流。",
            ],
            [
                "Harness",
                "任务大脑、满意度判定、任务状态机、升级调度、记忆沉淀。",
                "方向正确，但员工问卷没有直接表达满意度升级，它应作为底层支撑。",
            ],
            [
                "Hermes",
                "后台深度执行引擎，处理复杂任务、联网搜索、多步骤工具调用。",
                "适合作为 P1/P2 深度能力。当前员工更急需项目执行助理能力。",
            ],
            [
                "知识库与记忆",
                "解决方案太泛，注入公司资料、短期记忆、长期记忆。",
                "与问卷高度相关，但落地入口要从飞书文档/表格、客户项目资料开始。",
            ],
        ],
        [4.2, 10.4, 11.4],
    )

    start_page(doc, "五、需求与方案匹配度对比")
    add_table(
        doc,
        ["员工需求", "问卷强度", "Claude Code 覆盖情况", "判断与调整"],
        [
            [
                "任务提取和提醒",
                "91%",
                "部分覆盖",
                "Harness 有任务状态机基础，但需要明确开发任务提取、负责人、截止时间、提醒台账。",
            ],
            [
                "项目群聊总结",
                "82%",
                "覆盖不足",
                "Claude 方案强调深度执行，没有把群聊总结作为 P0 交付项，应补为二期第一模块。",
            ],
            [
                "飞书文档/表格资料查询",
                "82% / 73%",
                "方向匹配",
                "知识库方向正确，但落地应从飞书文档、飞书表格、客户资料白名单开始。",
            ],
            [
                "内容/方案生成",
                "55%-64%",
                "间接覆盖",
                "AstrBot LLM 能生成初稿，但需要按品宣和中台场景做模板。",
            ],
            [
                "低打扰交互",
                "高",
                "需要补充",
                "必须明确不默认群聊自动抢答，仅支持私聊、按钮、固定指令或 @ 触发。",
            ],
            [
                "Hermes 深度升级",
                "隐性需求",
                "高度覆盖",
                "适合用于方案太泛、资料不够、需要深挖场景，但不应抢在 P0 员工功能之前。",
            ],
        ],
        [4.3, 3.3, 5.2, 13.2],
    )

    start_page(doc, "六、最终调整建议：二期拆成 2A + 2B")
    add_callout(
        doc,
        "建议定版：DC-Agent 二期不推翻 Claude Code 方案，而是调整交付顺序。先做 2A 员工可感知能力，再做 2B 系统深度能力。这样既回应问卷，也不浪费已有架构设计。",
    )
    add_table(
        doc,
        ["阶段", "定位", "开发内容", "验收方式"],
        [
            [
                "2A",
                "员工可感知功能",
                "项目群聊总结、任务提取提醒、飞书资料查询、内容/方案生成模板、低打扰交互规则。",
                "员工能在飞书真实项目群和私聊里用起来，能减少重复整理和查资料时间。",
            ],
            [
                "2B",
                "系统深度能力",
                "Harness 满意度判定、Hermes 升级调度、知识库/记忆注入、多任务分流、失败重试。",
                "复杂任务能从 AstrBot 初稿升级为 Hermes 深度执行，并回传结果、沉淀记忆。",
            ],
        ],
        [2.8, 4.2, 9.8, 9.2],
    )

    start_page(doc, "七、建议时间轴")
    add_table(
        doc,
        ["时间", "阶段", "重点交付"],
        [
            [
                "2026-05-07 至 2026-05-10",
                "需求冻结与已有方案重排",
                "确认 DC-Agent 二期采用 2A + 2B 路线；把问卷最高频需求前置到 2A。",
            ],
            [
                "2026-05-11 至 2026-05-17",
                "2A-1 项目群聊总结",
                "完成客户反馈、老板要求、执行问题、阶段结论总结，支持项目群 @DC-Agent 触发。",
            ],
            [
                "2026-05-18 至 2026-05-24",
                "2A-2 任务提取和提醒",
                "完成负责人、事项、截止时间、待办状态识别，优先落地任务台账。",
            ],
            [
                "2026-05-25 至 2026-05-31",
                "2A-3 飞书资料查询",
                "接入首批飞书文档、飞书表格、客户资料、项目资料，回答标注来源。",
            ],
            [
                "2026-06-01 至 2026-06-07",
                "2A-4 内容与方案生成模板",
                "完成客户提案大纲、短视频脚本、口播文案、小红书/朋友圈文案、活动方案框架模板。",
            ],
            [
                "2026-06-08 至 2026-06-15",
                "灰度验收与 2B 接入评估",
                "选择 2 至 3 个真实项目群灰度，评估哪些任务进入 Harness 满意度判定和 Hermes 深度执行。",
            ],
            [
                "2026-06-16 起",
                "2B 系统深度能力推进",
                "推进 Harness 满意度检测、Hermes 升级调度、知识库/记忆注入和多任务分流。",
            ],
        ],
        [6.2, 5.4, 14.4],
    )

    start_page(doc, "八、定版结论")
    add_callout(
        doc,
        "最终建议：DC-Agent 二期应以“项目执行助手”为员工交付主题，以 Claude Code 已形成的 Harness/Hermes 架构为底层支撑。短期先交付项目群聊总结、任务提醒、资料查询和方案生成；中期再把不满意升级、深度执行、知识库记忆和多任务分流接起来。",
    )

    # Explicitly leave headers and footers blank.
    for section in doc.sections:
        section.header.is_linked_to_previous = False
        section.footer.is_linked_to_previous = False
        for paragraph in section.header.paragraphs:
            paragraph.text = ""
        for paragraph in section.footer.paragraphs:
            paragraph.text = ""

    OUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT_DOCX)


if __name__ == "__main__":
    build_report()
    print(OUT_DOCX)
